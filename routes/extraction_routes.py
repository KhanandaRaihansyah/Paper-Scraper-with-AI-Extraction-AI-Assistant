from flask import Blueprint, request, jsonify, send_file
import pandas as pd
import zipfile
import io
import time
import threading
import requests
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from models.models import SessionLocal, PaperExtraction
from services.llm_service import process_with_groq
from services.rangkuman_service import build_rangkuman_review, filter_records_by_titles

# Rate limiting: maksimal 1 ekstraksi LLM berjalan sekaligus.
# Mencegah ledakan panggilan API saat frontend mengirim banyak request serentak.
_extraction_semaphore = threading.Semaphore(1)

extraction_bp = Blueprint('extraction_routes', __name__)

@extraction_bp.route('/api/extract', methods=['POST'])
def extract_paper():
    # Rate limiting: antri masuk semaphore, timeout 10 detik.
    # Jika tidak bisa acquire (server sibuk), kembalikan 429 agar frontend retry.
    acquired = _extraction_semaphore.acquire(timeout=10)
    if not acquired:
        return jsonify({
            'error': 'Server sedang memproses paper lain. Silakan tunggu sebentar.',
            'retry': True
        }), 429

    try:
        data = request.json
        with SessionLocal() as db:
            # LANGKAH 1: Pencegahan Duplikat + Skip paper yang sebelumnya sudah gagal
            existing = db.query(PaperExtraction).filter_by(original_title=data['title']).first()

            if existing:
                if existing.extraction_status == 'failed':
                    return jsonify({
                        "message": "Paper sebelumnya gagal diekstrak dan dilewati",
                        "data": {"info": "skip_failed"}
                    })
                return jsonify({
                    "message": "Paper sudah pernah diekstrak sebelumnya",
                    "data": {"info": "skip"}
                })

            # LANGKAH 2: Panggil LLM (Hybrid: Groq → Gemini → Ollama)
            try:
                result_json = process_with_groq(
                    paper_title=data['title'],
                    abstract=data['abstract'],
                    authors=data['authors'],
                    year=data['year'],
                    context_keyword=data.get('keyword', 'teknologi'),
                    pdf_url=data.get('pdf_url')
                )
            except Exception as groq_error:
                # LLM gagal — catat sebagai 'failed' agar tidak dicoba ulang
                failed_record = PaperExtraction(
                    original_title=data['title'],
                    source_type=data.get('source'),
                    pdf_url=data.get('pdf_url'),
                    paper_url=data.get('paper_url'),
                    extraction_status='failed',
                    error_message=str(groq_error),
                )
                db.add(failed_record)
                db.commit()
                print(f"\n❌ Ekstraksi GAGAL untuk '{data['title'][:50]}': {groq_error}\n")
                return jsonify({"error": f"Ekstraksi gagal: {str(groq_error)}"}), 500

            # LANGKAH 3: Simpan hasil sukses ke Database
            extraction = PaperExtraction(
                original_title=data['title'],
                source_type=data['source'],
                search_keyword=data.get('keyword', 'teknologi'),
                pdf_url=data.get('pdf_url'),
                paper_url=data.get('paper_url'),
                extraction_status='success',
                doi=result_json.get('doi'),
                title=result_json.get('title'),
                authors=result_json.get('authors'),
                year=result_json.get('year'),
                abstract=result_json.get('abstract'),
                relevance=result_json.get('relevance'),
                systematic_review=result_json.get('systematic_review'),
                publisher=result_json.get('publisher'),
                application=result_json.get('application'),
                system=result_json.get('system'),
                algorithm=result_json.get('algorithm'),
                dataset=result_json.get('dataset'),
                keyword=result_json.get('keyword'),
                publication_type=result_json.get('publication_type'),
                journal_name=result_json.get('journal_name'),
                contribution=result_json.get('contribution'),
                limitations=result_json.get('limitations'),
            )
            db.add(extraction)
            db.commit()

        return jsonify({"message": "Berhasil diekstrak", "data": result_json})

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        _extraction_semaphore.release()

@extraction_bp.route('/api/extractions', methods=['GET'])
def get_extractions():
    with SessionLocal() as db:
        records = db.query(PaperExtraction).all()
        results = [
            {c.name: getattr(r, c.name) for c in r.__table__.columns}
            for r in records
        ]
    return jsonify({"extractions": results})

def _is_valid_field_value(value):
    """
    Helper untuk mengecek apakah nilai field dianggap valid (bukan kosong/tidak ada data).
    Return False jika nilai tidak valid (NULL, "Tidak diketahui", "-", dll.)
    """
    if not value:  # None atau string kosong
        return False
    value_lower = str(value).strip().lower()
    # Daftar kata kunci yang menandakan "tidak ada data"
    invalid_keywords = [
        'tidak diketahui', 'tidak disebutkan', 'tidak ada', 'tidak ditemukan',
        'unknown', 'not mentioned', 'not specified', 'n/a', 'na', '-', '–', '—'
    ]
    return value_lower not in invalid_keywords


def _should_exclude_paper(record):
    """
    Menentukan apakah paper harus diexclude dari export.
    
    Kriteria exclude:
    1. Systematic review = "Ya, ..." (literature review/survey)
    2. Keempat kolom (dataset, algorithm, system, application) semuanya tidak valid
    
    Return True jika paper harus diexclude.
    """
    # Kriteria 1: Exclude jika systematic review
    if record.systematic_review and str(record.systematic_review).strip().lower().startswith('ya,'):
        return True
    
    # Kriteria 2: Exclude jika SEMUA kolom teknis tidak valid
    has_valid_dataset = _is_valid_field_value(record.dataset)
    has_valid_algorithm = _is_valid_field_value(record.algorithm)
    has_valid_system = _is_valid_field_value(record.system)
    has_valid_application = _is_valid_field_value(record.application)
    
    # Jika SEMUA EMPAT kolom tidak valid, exclude
    if not (has_valid_dataset or has_valid_algorithm or has_valid_system or has_valid_application):
        return True
    
    return False


@extraction_bp.route('/api/extractions/export/<export_type>', methods=['GET'])
def export_extractions(export_type):
    keyword_filter = request.args.get('keyword', '').strip()
    db = SessionLocal()
    query = db.query(PaperExtraction)
    if keyword_filter and keyword_filter != 'all':
        query = query.filter((PaperExtraction.search_keyword == keyword_filter) | (PaperExtraction.keyword == keyword_filter))
    all_records = query.all()
    db.close()

    if not all_records:
        return jsonify({'error': 'Tidak ada data ekstraksi untuk diekspor'}), 404

    # FILTER: Exclude systematic review dan paper tanpa data teknis
    records = [r for r in all_records if not _should_exclude_paper(r)]
    
    if not records:
        return jsonify({'error': 'Tidak ada paper yang memenuhi kriteria export (semua paper adalah systematic review atau tidak memiliki data teknis)'}), 404

    if export_type in ['csv', 'excel']:
        data = [{c.name: getattr(r, c.name) for c in r.__table__.columns if c.name not in ['id', 'created_at']} for r in records]
        df = pd.DataFrame(data)
        
        output = io.BytesIO()
        if export_type == 'csv':
            df.to_csv(output, index=False, encoding='utf-8-sig')
            mimetype = 'text/csv'
            filename = 'hasil_ekstraksi_ai.csv'
        else:
            # Tulis ke Excel dengan openpyxl lalu aktifkan wrap text pada kolom teks panjang
            from openpyxl import load_workbook
            from openpyxl.styles import Alignment
            df.to_excel(output, index=False, engine='openpyxl')
            output.seek(0)
            wb = load_workbook(output)
            ws = wb.active
            # Kolom yang perlu wrap text: abstract, relevance, systematic_review,
            # contribution, limitations (cari posisi kolomnya dari header baris 1)
            wrap_cols = {'abstract', 'relevance', 'systematic_review', 'contribution', 'limitations'}
            col_indices = {}
            for cell in ws[1]:
                if str(cell.value).lower() in wrap_cols:
                    col_indices[cell.column] = str(cell.value)
            # Terapkan wrap text ke semua sel di kolom tersebut
            for row in ws.iter_rows(min_row=2):
                for cell in row:
                    if cell.column in col_indices:
                        cell.alignment = Alignment(wrap_text=True, vertical='top')
            # Set lebar kolom abstract sedikit lebih lebar agar mudah dibaca
            for col_num, col_name in col_indices.items():
                col_letter = ws.cell(row=1, column=col_num).column_letter
                ws.column_dimensions[col_letter].width = 60
            output = io.BytesIO()
            wb.save(output)
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            filename = 'hasil_ekstraksi_ai.xlsx'
            
        output.seek(0)
        return send_file(output, as_attachment=True, download_name=filename, mimetype=mimetype)
        
    elif export_type == 'word':
        doc = Document()

        # --- Setup margin halaman (2.5 cm semua sisi) ---
        section = doc.sections[0]
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        # --- Helper: set lebar kolom tabel (dalam twips, 1 cm = 567 twips) ---
        def _set_col_width(cell, width_cm):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(width_cm * 567)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        # --- Helper: hilangkan semua border tabel agar terlihat bersih ---
        def _remove_table_borders(table):
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)
            tblPr.append(tblBorders)

        # --- Helper: tambah teks ke cell dengan font size tertentu ---
        def _cell_text(cell, text, font_size_pt=11, bold=False):
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(2)
            run = para.add_run(str(text) if text else '-')
            run.font.size = Pt(font_size_pt)
            run.bold = bold
            return para

        # Lebar teks area = lebar kertas A4 (21 cm) - margin kiri (2.5) - margin kanan (2.5) = 16 cm
        # Kolom label: 5 cm | kolom titik dua: 0.6 cm | kolom nilai: 10.4 cm
        COL_LABEL = 5.0
        COL_COLON = 0.6
        COL_VALUE = 10.4

        # --- Judul dokumen: DETAIL REVIEW (bold, centered) ---
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(12)
        title_run = title_para.add_run('DETAIL REVIEW')
        title_run.bold = True
        title_run.font.size = Pt(12)

        # --- Iterasi setiap paper ---
        for idx, r in enumerate(records):
            fields = [
                ('title',                  r.title),
                ('abstract',               r.abstract),
                ('authors',                r.authors),
                ('year',                   str(r.year) if r.year else '-'),
                ('journal link',           r.paper_url),
                ('contribution / novelty', r.contribution),
                ('limitations',            r.limitations),
            ]

            # Buat tabel 7 baris x 3 kolom (label | : | nilai)
            table = doc.add_table(rows=len(fields), cols=3)
            _remove_table_borders(table)

            for row_idx, (label, value) in enumerate(fields):
                row = table.rows[row_idx]

                # Kolom 0: label
                _set_col_width(row.cells[0], COL_LABEL)
                _cell_text(row.cells[0], label)

                # Kolom 1: titik dua
                _set_col_width(row.cells[1], COL_COLON)
                _cell_text(row.cells[1], ':')

                # Kolom 2: nilai (teks panjang akan wrap di dalam sel ini saja)
                _set_col_width(row.cells[2], COL_VALUE)
                _cell_text(row.cells[2], value or '-')

            # Spasi antar paper (paragraph kosong setelah tabel)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after  = Pt(6)

        # Simpan ke BytesIO
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            as_attachment=True,
            download_name='detail_review.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    elif export_type == 'pdf-zip':
        memory_file = io.BytesIO()
        with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
            for i, r in enumerate(records):
                if r.pdf_url:
                    try:
                        response = requests.get(r.pdf_url, timeout=10)
                        if response.status_code == 200:
                            safe_title = "".join([c for c in r.title if c.isalpha() or c.isdigit() or c==' ']).rstrip()[:50]
                            zf.writestr(f"{i+1}_{safe_title}.pdf", response.content)
                            time.sleep(5) # Aturan jeda PRD
                    except Exception:
                        continue
        memory_file.seek(0)
        return send_file(memory_file, as_attachment=True, download_name='kumpulan_pdf_ekstraksi.zip', mimetype='application/zip')


@extraction_bp.route('/api/extractions/export/rangkuman-review', methods=['GET'])
def export_rangkuman_review():
    """
    Menghasilkan dokumen Word 'Rangkuman Review' berisi:
      1. Uraian — grafik rekapitulasi (per tahun, per sumber, per tipe publikasi)
      2. Kontribusi — rangkuman contribution dari setiap paper
      3. Keterbatasan — rangkuman limitations dari setiap paper
    Hanya paper dengan extraction_status='success' yang diikutsertakan.
    """
    try:
        keyword_filter = request.args.get('keyword', '').strip()
        with SessionLocal() as db:
            query = db.query(PaperExtraction).filter_by(extraction_status='success')
            if keyword_filter and keyword_filter != 'all':
                query = query.filter((PaperExtraction.search_keyword == keyword_filter) | (PaperExtraction.keyword == keyword_filter))
            all_records = query.all()
            if not all_records:
                return jsonify({'error': 'Tidak ada data ekstraksi yang tersedia'}), 404
            output = build_rangkuman_review(all_records)

        return send_file(
            output,
            as_attachment=True,
            download_name='rangkuman_review.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
