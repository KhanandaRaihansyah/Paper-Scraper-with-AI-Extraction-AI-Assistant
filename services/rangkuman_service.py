"""
rangkuman_service.py
Menghasilkan dokumen Word "Rangkuman Review" berisi:
  1. Uraian  — grafik rekapitulasi + narasi ringkasan keseluruhan
  2. Kontribusi — rangkuman tematik kontribusi/novelty dari seluruh paper
  3. Keterbatasan — rangkuman tematik keterbatasan dari seluruh paper
"""

import io
import re
from collections import Counter

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# Helpers dokumen
# ─────────────────────────────────────────────

def _set_margins(doc, cm_val=2.5):
    section = doc.sections[0]
    for attr in ('left_margin', 'right_margin', 'top_margin', 'bottom_margin'):
        setattr(section, attr, Cm(cm_val))


def _heading(doc, text, bold=True, size_pt=13, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=Pt(12), space_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def _body(doc, text, size_pt=11, space_after=Pt(6), indent_cm=0, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(str(text) if text else '-')
    run.font.size = Pt(size_pt)
    run.bold = bold
    return p


def _bullet(doc, text, size_pt=11, indent_cm=0.7):
    """Paragraf dengan bullet •"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(str(text) if text else '-')
    run.font.size = Pt(size_pt)
    return p


# ─────────────────────────────────────────────
# Normalisasi tipe publikasi
# ─────────────────────────────────────────────

_PUBTYPE_MAP = {
    'jurnal': 'Jurnal', 'journal': 'Jurnal', 'journal article': 'Jurnal',
    'review article': 'Jurnal', 'artikel ilmiah': 'Jurnal', 'artikel': 'Jurnal',
    'article': 'Jurnal', 'conference paper': 'Konferensi', 'conference': 'Konferensi',
    'konferensi': 'Konferensi', 'short paper': 'Konferensi', 'preprint': 'Preprint',
    'survei': 'Survei', 'survey': 'Survei', 'survey paper': 'Survei',
    'bibliometric survey': 'Survei', 'tinjauan pustaka': 'Survei', 'review': 'Survei',
    'white paper': 'Lainnya', 'technical report': 'Lainnya',
    'laporan sistem': 'Lainnya', 'dataset': 'Lainnya', 'position paper': 'Lainnya',
    'tidak diketahui': 'Tidak Diketahui',
}

def _normalize_pubtype(raw):
    if not raw:
        return 'Tidak Diketahui'
    return _PUBTYPE_MAP.get(raw.strip().lower(), raw.strip())


# ─────────────────────────────────────────────
# Grafik
# ─────────────────────────────────────────────

CHART_COLORS = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#3B1F2B',
                '#44BBA4', '#E94F37', '#393E41', '#F5A623', '#7B2D8B']


def _fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_per_tahun(records):
    year_counts = Counter(str(r.year) for r in records if r.year and str(r.year).isdigit())
    years = sorted(year_counts.keys())
    counts = [year_counts[y] for y in years]
    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.bar(years, counts, color=CHART_COLORS[:len(years)], edgecolor='white', linewidth=0.8)
    ax.set_title('Distribusi Paper per Tahun Publikasi', fontsize=13, fontweight='bold', pad=12)
    ax.set_xlabel('Tahun', fontsize=11)
    ax.set_ylabel('Jumlah Paper', fontsize=11)
    ax.set_xticks(range(len(years)))
    ax.set_xticklabels(years, rotation=0)
    ax.yaxis.get_major_locator().set_params(integer=True)
    for bar, count in zip(bars, counts):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2,
                str(count), ha='center', va='bottom', fontsize=10, fontweight='bold')
    ax.spines[['top', 'right']].set_visible(False)
    ax.set_facecolor('#F8F9FA')
    fig.patch.set_facecolor('white')
    fig.tight_layout()
    return _fig_to_bytes(fig), year_counts, years, counts


def _chart_per_sumber(records):
    source_counts = Counter(r.source_type for r in records if r.source_type)
    labels = list(source_counts.keys())
    sizes = [source_counts[l] for l in labels]
    total = sum(sizes)
    fig, ax = plt.subplots(figsize=(6, 5))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=None, colors=CHART_COLORS[:len(labels)],
        autopct=lambda p: f'{p:.1f}%\n({int(round(p * total / 100))})',
        startangle=140, pctdistance=0.75,
        wedgeprops={'edgecolor': 'white', 'linewidth': 1.5}
    )
    for at in autotexts:
        at.set_fontsize(10)
        at.set_fontweight('bold')
    legend_patches = [mpatches.Patch(color=CHART_COLORS[i], label=f'{labels[i]} ({sizes[i]})')
                      for i in range(len(labels))]
    ax.legend(handles=legend_patches, loc='lower center', bbox_to_anchor=(0.5, -0.12),
              ncol=2, fontsize=10, frameon=False)
    ax.set_title('Distribusi Paper per Sumber Data', fontsize=13, fontweight='bold', pad=12)
    fig.tight_layout()
    return _fig_to_bytes(fig), source_counts


def _chart_per_pubtype(records):
    raw_counts = Counter(_normalize_pubtype(r.publication_type) for r in records)
    sorted_items = sorted(raw_counts.items(), key=lambda x: x[1], reverse=True)
    labels = [it[0] for it in sorted_items]
    counts = [it[1] for it in sorted_items]
    fig, ax = plt.subplots(figsize=(8, max(3, len(labels) * 0.65)))
    bars = ax.barh(labels[::-1], counts[::-1], color=CHART_COLORS[:len(labels)],
                   edgecolor='white', linewidth=0.8)
    ax.set_title('Distribusi Paper per Tipe Publikasi', fontsize=13, fontweight='bold', pad=12)
    ax.set_xlabel('Jumlah Paper', fontsize=11)
    ax.xaxis.get_major_locator().set_params(integer=True)
    for bar, count in zip(bars, counts[::-1]):
        ax.text(bar.get_width() + 0.15, bar.get_y() + bar.get_height() / 2,
                str(count), va='center', fontsize=10, fontweight='bold')
    ax.spines[['top', 'right']].set_visible(False)
    ax.set_facecolor('#F8F9FA')
    fig.patch.set_facecolor('white')
    fig.tight_layout()
    return _fig_to_bytes(fig), raw_counts


# ─────────────────────────────────────────────
# Pengkategorian tematik kontribusi & keterbatasan
# (Dinamis: tema dibangun dari keyword paper, bukan hardcoded)
# ─────────────────────────────────────────────

import re

def _extract_top_keywords(records, top_n=10):
    """
    Mengekstrak kata kunci dominan dari seluruh kumpulan paper.
    Sumber: field keyword, algorithm, system, application.
    Mengembalikan list string kata kunci unik, diurutkan berdasarkan frekuensi.
    """
    # Stop words yang diabaikan
    _stopwords = {
        'the', 'a', 'an', 'and', 'or', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by',
        'is', 'are', 'was', 'were', 'based', 'using', 'via', 'this', 'that', 'from',
        'yang', 'dan', 'di', 'ke', 'dari', 'pada', 'dengan', 'untuk', 'secara', 'atau',
        'ini', 'itu', 'adalah', 'sebagai', 'dalam', 'oleh', 'bisa', 'dapat', 'serta',
        'tidak', 'paper', 'penelitian', 'study', 'approach', 'method', 'proposed'
    }
    freq = Counter()
    for r in records:
        sources = ' '.join(filter(None, [
            r.keyword or '', r.algorithm or '',
            r.system or '', r.application or ''
        ]))
        tokens = re.sub(r'[^\w\s/-]', ' ', sources.lower()).split()
        for tok in tokens:
            tok = tok.strip('/-_')
            if len(tok) >= 3 and tok not in _stopwords:
                freq[tok] += 1
    # Ambil top_n kata kunci paling sering muncul
    return [kw for kw, _ in freq.most_common(top_n)]


def _build_dynamic_themes(records):
    """
    Membangun daftar tema dinamis dari keyword paper yang ada.
    Setiap tema berupa tuple (regex_pattern, label_tema).

    Strategi:
    - Ambil top-10 keyword dari seluruh paper.
    - Kelompokkan keyword yang serupa ke dalam satu tema menggunakan prefix match.
    - Tambahkan tema fallback generik (dataset, skala, simulasi) yang selalu berguna.
    """
    top_kws = _extract_top_keywords(records, top_n=12)

    # Buat tema dari kata kunci teratas — setiap kata kunci menjadi pola regex sendiri
    dynamic_themes = []
    seen_labels = set()
    for kw in top_kws:
        label = kw.replace('-', ' ').replace('/', ' ').title()
        if label not in seen_labels:
            seen_labels.add(label)
            # Pola: cocokkan kata kunci itu sendiri (bisa gabungan 2-3 kata)
            pattern = re.escape(kw)
            dynamic_themes.append((pattern, label))

    return dynamic_themes


def _build_dynamic_limit_themes():
    """
    Tema keterbatasan: selalu relevan untuk penelitian apapun.
    Tidak di-hardcode ke domain IoT, tapi menggunakan pola universal.
    """
    return [
        (r'dataset|data yang terbatas|data kecil|jumlah sampel|representatif|real.?world',
         'Keterbatasan Dataset dan Validasi Empiris'),
        (r'simulasi|tidak diuji|kondisi nyata|pengujian terbatas|skenario terbatas|laborator',
         'Ketergantungan pada Simulasi dan Pengujian Terbatas'),
        (r'keamanan|privasi|serangan|enkripsi|rentan|vulnerability|attack',
         'Aspek Keamanan dan Privasi yang Belum Dieksplorasi'),
        (r'energi|daya|resource|memori|komputasi|efisiensi|overhead|latency',
         'Keterbatasan Sumber Daya dan Efisiensi'),
        (r'skala|besar|distribusi|heterogen|kompleks|deployment|generali',
         'Skalabilitas dan Generalisasi'),
        (r'ruang lingkup|sempit|spesifik|hanya|satu domain|tidak membahas|terbatas pada',
         'Ruang Lingkup Penelitian yang Sempit'),
        (r'biaya|cost|implementasi|produksi|industri|adopsi|praktis',
         'Tantangan Implementasi Praktis'),
    ]


def _normalize_title(t):
    if not t:
        return ""
    t = t.lower().strip()
    t = re.sub(r'[^\w\s]', ' ', t)
    t = re.sub(r'\s+', ' ', t).strip()
    return t


def filter_records_by_titles(records, title_list):
    """
    Kembalikan subset records yang judulnya (title atau original_title)
    cocok dengan salah satu judul di title_list.
    Pencocokan case-insensitive dan toleran terhadap tanda baca.
    """
    norm_targets = {_normalize_title(t) for t in title_list}

    result = []
    for r in records:
        key = _normalize_title(r.title or r.original_title or '')
        if key in norm_targets:
            result.append(r)
            continue
        # Partial match fallback (50 karakter pertama)
        if any(key[:50] in t or t[:50] in key for t in norm_targets if len(t) > 20):
            result.append(r)
    return result


def _classify_themes(text, theme_list):
    """Kembalikan label tema pertama yang cocok dengan teks, atau 'Lainnya'."""
    if not text:
        return None
    text_lower = text.lower()
    for pattern, label in theme_list:
        if re.search(pattern, text_lower):
            return label
    return 'Lainnya'


def _group_by_theme(records, field, theme_list):
    """
    Kembalikan dict: {tema: [teks_field, ...]}
    field: 'contribution' atau 'limitations'
    """
    groups = {}
    _invalid = {'tidak diketahui', 'tidak disebutkan', 'tidak ada', '-', ''}
    for r in records:
        val = getattr(r, field, None)
        if not val or str(val).strip().lower() in _invalid:
            continue
        theme = _classify_themes(str(val), theme_list)
        if theme:
            groups.setdefault(theme, []).append(str(val).strip())
    return groups


# ─────────────────────────────────────────────
# Narasi ringkasan tematik
# Mengubah list teks mentah → 1 paragraf narasi yang kohesif
# ─────────────────────────────────────────────

def _build_contrib_narrative(theme, texts):
    """
    Menghasilkan satu paragraf narasi kontribusi per tema
    berdasarkan pola teks yang dikumpulkan.
    """
    count = len(texts)

    # Ekstrak kata kunci unik dari teks (pendekatan heuristic sederhana)
    # Ambil semua frasa setelah "mencakup:" atau setelah ":" pertama
    key_points = []
    for t in texts:
        # Ambil bagian setelah "mencakup:" jika ada
        m = re.search(r'mencakup[:\s]+(.+)', t, re.IGNORECASE | re.DOTALL)
        if m:
            point = m.group(1).strip()
        else:
            # Gunakan teks apa adanya, batasi 120 karakter
            point = t[:120]
        # Bersihkan nomor di awal seperti "(1)", "1."
        point = re.sub(r'^\s*[\(\d]+[\)\.]\s*', '', point)
        point = point.split(';')[0].split(',')[0].strip()
        if len(point) > 20:
            key_points.append(point[:100])

    # Ambil maksimal 4 poin unik representatif
    seen = set()
    unique_points = []
    for p in key_points:
        key = p[:40].lower()
        if key not in seen:
            seen.add(key)
            unique_points.append(p)
        if len(unique_points) >= 4:
            break

    if not unique_points:
        return (f'Sebanyak {count} penelitian dalam tema ini memberikan kontribusi '
                f'signifikan terhadap pengembangan {theme.lower()}.')

    intro = f'Dari {count} penelitian yang termasuk dalam tema {theme}, terdapat beberapa kontribusi utama yang menonjol. '
    if len(unique_points) == 1:
        body = f'Kontribusi yang paling dominan adalah {unique_points[0].lower()}.'
    else:
        joined = '; '.join(p.rstrip('.').lower() for p in unique_points[:-1])
        last = unique_points[-1].rstrip('.').lower()
        body = f'Di antaranya mencakup {joined}, serta {last}.'
    return intro + body


def _build_limit_narrative(theme, texts):
    """Menghasilkan satu paragraf narasi keterbatasan per tema."""
    count = len(texts)
    key_points = []
    for t in texts:
        m = re.search(r'mencakup[:\s]+(.+)', t, re.IGNORECASE | re.DOTALL)
        if m:
            point = m.group(1).strip()
        else:
            point = t[:120]
        point = re.sub(r'^\s*[\(\d]+[\)\.]\s*', '', point)
        point = point.split(';')[0].split(',')[0].strip()
        if len(point) > 20:
            key_points.append(point[:100])

    seen = set()
    unique_points = []
    for p in key_points:
        key = p[:40].lower()
        if key not in seen:
            seen.add(key)
            unique_points.append(p)
        if len(unique_points) >= 4:
            break

    if not unique_points:
        return (f'Sebanyak {count} penelitian dalam tema {theme.lower()} '
                f'mengidentifikasi keterbatasan yang perlu diatasi pada penelitian mendatang.')

    intro = f'Pada tema {theme}, sebanyak {count} penelitian mengungkap keterbatasan yang serupa. '
    if len(unique_points) == 1:
        body = f'Keterbatasan utama yang ditemukan adalah {unique_points[0].lower()}.'
    else:
        joined = '; '.join(p.rstrip('.').lower() for p in unique_points[:-1])
        last = unique_points[-1].rstrip('.').lower()
        body = f'Keterbatasan yang paling umum mencakup {joined}, serta {last}.'
    return intro + body


# ─────────────────────────────────────────────
# Builder utama
# ─────────────────────────────────────────────

def build_rangkuman_review(records):
    """
    Menerima list PaperExtraction (sudah difilter success).
    Mengembalikan BytesIO berisi file .docx.

    Struktur:
      1. Uraian — grafik rekapitulasi + narasi ringkasan
      2. Kontribusi — rangkuman tematik (bukan per paper)
      3. Keterbatasan — rangkuman tematik (bukan per paper)
    """
    doc = Document()
    _set_margins(doc, 2.5)
    total = len(records)

    # ── Judul ──────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(20)
    title_p.paragraph_format.space_before = Pt(0)
    title_run = title_p.add_run('RANGKUMAN REVIEW')
    title_run.bold = True
    title_run.font.size = Pt(16)

    # ══════════════════════════════════════════════════════════════
    # SECTION 1 — URAIAN
    # ══════════════════════════════════════════════════════════════
    _heading(doc, '1.   Uraian', size_pt=12, space_before=Pt(4))

    # Hitung statistik untuk narasi
    year_counts_raw = Counter(str(r.year) for r in records if r.year and str(r.year).isdigit())
    years_sorted = sorted(year_counts_raw.keys())
    source_counts_raw = Counter(r.source_type for r in records if r.source_type)
    pubtype_norm = Counter(_normalize_pubtype(r.publication_type) for r in records)
    pubtype_top3 = pubtype_norm.most_common(3)

    year_max = max(year_counts_raw, key=year_counts_raw.get) if year_counts_raw else '-'
    year_min_val = min(year_counts_raw, key=year_counts_raw.get) if year_counts_raw else '-'
    src_dominant = source_counts_raw.most_common(1)[0] if source_counts_raw else ('-', 0)
    pubtype_desc = ', '.join(f'{k} ({v} paper)' for k, v in pubtype_top3)

    narasi_uraian = (
        f'Tinjauan literatur ini mencakup {total} artikel ilmiah yang telah dianalisis '
        f'menggunakan sistem ekstraksi berbasis AI. Rentang publikasi mencakup tahun '
        f'{min(years_sorted) if years_sorted else "-"} hingga {max(years_sorted) if years_sorted else "-"}, '
        f'dengan konsentrasi tertinggi pada tahun {year_max} ({year_counts_raw.get(year_max, 0)} paper). '
        f'Sumber pengumpulan data didominasi oleh {src_dominant[0]} '
        f'({src_dominant[1]} paper, {src_dominant[1]/total*100:.1f}% dari total). '
        f'Berdasarkan tipe publikasi, mayoritas artikel termasuk dalam kategori {pubtype_desc}. '
        f'Keragaman topik yang dikaji mencerminkan luasnya cakupan riset dalam bidang '
        f'Internet of Things, sistem tertanam, dan teknologi terkait.'
    )
    _body(doc, narasi_uraian, size_pt=11, space_after=Pt(10))


    # — Grafik 1: per Tahun ─────────────────────────────────────────
    _heading(doc, 'a.   Distribusi Paper per Tahun Publikasi',
             size_pt=11, bold=False, space_before=Pt(6), space_after=Pt(4))

    img_tahun, year_counts, years, counts = _chart_per_tahun(records)
    doc.add_picture(img_tahun, width=Cm(14))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    trend_desc = 'cenderung meningkat' if counts[-1] >= counts[0] else 'cenderung menurun'
    _body(doc,
          f'Gambar 1 menunjukkan distribusi paper berdasarkan tahun publikasi. '
          f'Publikasi {trend_desc} selama periode {min(years)}–{max(years)}, '
          f'dengan puncak pada tahun {year_max} sebanyak {year_counts.get(year_max, 0)} paper. '
          f'Tren ini menggambarkan dinamika perkembangan riset di bidang yang dikaji.',
          size_pt=10.5, space_after=Pt(8))

    # — Grafik 2: per Sumber ────────────────────────────────────────
    _heading(doc, 'b.   Distribusi Paper per Sumber Data',
             size_pt=11, bold=False, space_before=Pt(6), space_after=Pt(4))

    img_sumber, source_counts = _chart_per_sumber(records)
    doc.add_picture(img_sumber, width=Cm(10))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    src_list = ', '.join(f'{k} ({v} paper)' for k, v in source_counts.most_common())
    _body(doc,
          f'Gambar 2 memperlihatkan komposisi sumber data yang digunakan. '
          f'Paper diperoleh dari {len(source_counts)} sumber, yakni {src_list}. '
          f'Dominasi {src_dominant[0]} menunjukkan bahwa repositori tersebut merupakan '
          f'sumber utama literatur dalam topik ini.',
          size_pt=10.5, space_after=Pt(8))

    # — Grafik 3: per Tipe Publikasi ────────────────────────────────
    _heading(doc, 'c.   Distribusi Paper per Tipe Publikasi',
             size_pt=11, bold=False, space_before=Pt(6), space_after=Pt(4))

    img_pubtype, pubtype_counts = _chart_per_pubtype(records)
    doc.add_picture(img_pubtype, width=Cm(14))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    pt_top = pubtype_counts.most_common(1)[0]
    _body(doc,
          f'Gambar 3 menyajikan sebaran paper berdasarkan tipe publikasi. '
          f'Tipe yang paling banyak ditemukan adalah {pt_top[0]} ({pt_top[1]} paper), '
          f'diikuti oleh {", ".join(f"{k} ({v})" for k, v in pubtype_counts.most_common()[1:3])}. '
          f'Keberagaman ini mencerminkan cakupan literatur yang komprehensif.',
          size_pt=10.5, space_after=Pt(8))


    # ══════════════════════════════════════════════════════════════
    # SECTION 2 — KONTRIBUSI (rangkuman tematik)
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, '2.   Kontribusi (Contributions)', size_pt=12, space_before=Pt(4))

    _body(doc,
          f'Berdasarkan analisis terhadap {total} artikel, kontribusi dan kebaruan '
          f'penelitian dapat dikelompokkan ke dalam beberapa tema utama sebagai berikut.',
          size_pt=11, space_after=Pt(8))

    # Bangun tema secara dinamis dari keyword paper yang ada
    contrib_themes = _build_dynamic_themes(records)
    contrib_groups = _group_by_theme(records, 'contribution', contrib_themes)

    # Urutkan dari tema dengan jumlah terbanyak
    sorted_contrib = sorted(contrib_groups.items(), key=lambda x: len(x[1]), reverse=True)

    for idx, (theme, texts) in enumerate(sorted_contrib, 1):
        # Sub-heading tema
        _heading(doc, f'{idx}.   {theme}',
                 size_pt=11, bold=True,
                 color=(0x1A, 0x53, 0x76),
                 space_before=Pt(10), space_after=Pt(3))

        # Narasi ringkasan
        narrative = _build_contrib_narrative(theme, texts)
        _body(doc, narrative, size_pt=11, space_after=Pt(6), indent_cm=0.5)

    # Penutup section 2
    _body(doc,
          f'Secara keseluruhan, penelitian-penelitian yang dikaji telah memberikan '
          f'kontribusi nyata dalam memperluas batas pengetahuan dalam bidang yang dikaji. '
          f'Inovasi yang dihasilkan mencerminkan arah riset yang semakin multidisiplin '
          f'dan berorientasi pada implementasi dunia nyata.',
          size_pt=11, space_after=Pt(8))


    # ══════════════════════════════════════════════════════════════
    # SECTION 3 — KETERBATASAN (rangkuman tematik)
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, '3.   Keterbatasan (Limitations)', size_pt=12, space_before=Pt(4))

    _body(doc,
          f'Analisis terhadap keterbatasan dari {total} artikel mengungkap beberapa '
          f'pola berulang yang perlu mendapat perhatian pada penelitian mendatang. '
          f'Berikut adalah rangkuman berdasarkan tema keterbatasan yang paling umum ditemukan.',
          size_pt=11, space_after=Pt(8))

    limit_themes = _build_dynamic_limit_themes()
    limit_groups = _group_by_theme(records, 'limitations', limit_themes)
    sorted_limits = sorted(limit_groups.items(), key=lambda x: len(x[1]), reverse=True)

    for idx, (theme, texts) in enumerate(sorted_limits, 1):
        _heading(doc, f'{idx}.   {theme}',
                 size_pt=11, bold=True,
                 color=(0xC0, 0x39, 0x2B),
                 space_before=Pt(10), space_after=Pt(3))

        narrative = _build_limit_narrative(theme, texts)
        _body(doc, narrative, size_pt=11, space_after=Pt(6), indent_cm=0.5)

    # Penutup section 3
    _body(doc,
          f'Keterbatasan-keterbatasan tersebut membuka peluang riset lanjutan yang '
          f'signifikan, khususnya dalam hal pengujian pada kondisi nyata dengan dataset '
          f'yang lebih besar dan representatif, serta peningkatan aspek keamanan dan '
          f'generalisasi lintas domain yang lebih luas. Penelitian selanjutnya '
          f'diharapkan dapat mengatasi gap tersebut untuk menghasilkan solusi '
          f'yang lebih robust, efisien, dan siap pakai.',
          size_pt=11, space_after=Pt(8))

    # ─── Simpan ke BytesIO ─────────────────────────────────────────
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
